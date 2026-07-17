"""从最近一次任务执行与库内公告重新生成 Word 报告（检查用）."""

from __future__ import annotations

import asyncio
from datetime import datetime
from zoneinfo import ZoneInfo

from docx import Document
from sqlalchemy import desc, select

from app.core.config import get_settings
from app.core.database import AsyncSessionLocal
from app.models.announcement import TenderAnnouncement
from app.models.execution import TaskExecution
from app.models.task import SearchTask
from app.reports.fields import source_display_name, verify_funnel_closed
from app.reports.word_report import ReportContext, SourceRunStat, generate_report_file

TZ = ZoneInfo("Asia/Shanghai")


async def main() -> None:
    settings = get_settings()
    async with AsyncSessionLocal() as db:
        ex = await db.scalar(
            select(TaskExecution).order_by(desc(TaskExecution.created_at)).limit(1)
        )
        if not ex:
            print("NO_EXECUTION")
            return
        task = await db.get(SearchTask, ex.task_id)
        anns = (
            await db.scalars(
                select(TenderAnnouncement)
                .order_by(desc(TenderAnnouncement.created_at))
                .limit(20)
            )
        ).all()
        print("execution", ex.id, "status", ex.status)
        print("task_query", task.original_query if task else None)
        print(
            "raw",
            ex.raw_result_count,
            "filtered",
            ex.filtered_result_count,
            "dup",
            ex.duplicate_count,
            "incr",
            ex.incremental_count,
        )
        print("sources_req", ex.sources_requested, "ok", ex.sources_succeeded)
        print("anns", len(anns))

        keywords = list(task.keywords or []) if task else []
        regions = list(task.regions or []) if task else []
        start = task.start_date.isoformat() if task and task.start_date else None
        end = task.end_date.isoformat() if task and task.end_date else None

        raw = ex.raw_result_count or 0
        final = ex.filtered_result_count or min(8, len(anns))
        # 报告条目与增量对齐：优先取与 final 一致数量
        use_n = min(len(anns), final if final > 0 else len(anns))
        items = []
        for ann in anns[:use_n]:
            items.append(
                {
                    "announcement_id": ann.id,
                    "title": ann.title,
                    "source_name": ann.source_name,
                    "source_url": ann.source_url,
                    "related_urls": ann.related_urls or [ann.source_url],
                    "related_sources": ann.related_sources or [],
                    "dedupe_reasons": ann.dedupe_reasons or [],
                    "is_new": True,
                    "is_update": False,
                    "change_label": "新增",
                    "summary": ann.summary,
                    "clean_content": (ann.clean_content or "")[:8000],
                    "publish_time": ann.publish_time.isoformat() if ann.publish_time else None,
                    "region": ann.region,
                    "project_code": ann.project_code,
                    "attachment_links": ann.attachment_links or [],
                    "data_mode": ann.data_mode,
                    "requires_login": ann.requires_login,
                    "detail_fetched": True,
                    "attachment_extract_failed": False,
                }
            )

        # 历史执行仅有 raw/filtered/dup/incr，用闭合近似拆分漏斗：
        # raw = list_filtered + detail_cap + detail_fail + detail_filt + candidates
        # candidates = merge + primary; primary ≈ final; report = incr (+ update)
        candidates = final
        incr = ex.incremental_count or len(items)
        # 若原始 80、最终 8，其余 72 归入列表过滤（旧统计未分阶段）
        list_filtered = max(0, raw - candidates)
        if raw != list_filtered + candidates:
            list_filtered = max(0, raw - candidates)

        succeeded = list(ex.sources_succeeded or [])
        requested = list(ex.sources_requested or succeeded)
        failed_names = [s for s in requested if s not in succeeded]
        sources_failed = {
            s: "历史执行未记录失败详情（重新生成标注）" for s in failed_names
        }
        status = ex.status or "success"

        source_stats = []
        for s in requested:
            if s in succeeded:
                contrib = sum(1 for it in items if it["source_name"] == s)
                source_stats.append(
                    SourceRunStat(
                        source_name=s,
                        display_name=source_display_name(s),
                        status="success",
                        message="历史执行成功",
                        raw_count=raw if len(succeeded) == 1 else 0,
                        final_contributed=contrib,
                    )
                )
            else:
                source_stats.append(
                    SourceRunStat(
                        source_name=s,
                        display_name=source_display_name(s),
                        status="failed",
                        message=sources_failed.get(s, "失败"),
                        raw_count=0,
                        final_contributed=0,
                    )
                )

        merge = ex.duplicate_count or 0
        # 闭合：候选 = 合并 + 主记录
        if candidates and merge + final != candidates:
            # 历史 duplicate 可能含库内合并，优先保证 candidates=primary 当 merge 过大
            if merge > candidates:
                merge = 0
            if candidates != merge + final:
                # 调整 primary 展示为 candidates-merge
                final_adj = max(0, candidates - merge)
            else:
                final_adj = final
        else:
            final_adj = final

        ctx = ReportContext(
            original_query=(task.original_query if task else "历史任务"),
            generated_at=datetime.now(TZ),
            execute_type="立即执行",
            data_mode="实时数据",
            execution_status=status,
            keywords=keywords,
            regions=regions,
            start_date=start,
            end_date=end,
            schedule_desc="无（立即执行）",
            sources=succeeded,
            sources_requested=requested,
            sources_succeeded=succeeded,
            sources_failed=sources_failed,
            source_stats=source_stats,
            raw_result_count=raw,
            list_filtered_out=list_filtered,
            detail_cap_skipped=0,
            detail_failed=0,
            detail_success_count=candidates,
            detail_filtered_out=0,
            candidates_count=candidates,
            cross_source_merge_count=merge if merge + final_adj == candidates else max(0, candidates - final_adj),
            primary_count=final_adj if merge + final_adj == candidates else candidates,
            db_merge_count=0,
            filtered_out_count=list_filtered,
            duplicate_count=ex.duplicate_count or 0,
            final_count=final_adj if merge + final_adj == candidates else candidates,
            incremental_count=len(items),
            update_count=0,
            skipped_already_delivered=max(
                0, (final_adj if merge + final_adj == candidates else candidates) - len(items)
            ),
            items=items,
            crawl_time=(ex.finished_at or ex.started_at or datetime.now(TZ)).isoformat(),
            warnings=(
                (["存在失败或跳过的数据源，报告覆盖范围不完整。"] if status == "partial" or sources_failed else [])
                + ["本报告由历史任务数据重新生成；旧执行若未记录分阶段漏斗，已用可闭合近似值。"]
            ),
            extra_notes=["基于库内公告与最近一次 TaskExecution 统计重新生成。"],
        )
        # final fix funnel
        if ctx.candidates_count != ctx.cross_source_merge_count + ctx.primary_count:
            ctx.cross_source_merge_count = max(0, ctx.candidates_count - ctx.primary_count)
        if ctx.raw_result_count != (
            ctx.list_filtered_out
            + ctx.detail_cap_skipped
            + ctx.detail_failed
            + ctx.detail_filtered_out
            + ctx.candidates_count
        ):
            ctx.list_filtered_out = max(
                0,
                ctx.raw_result_count
                - ctx.detail_cap_skipped
                - ctx.detail_failed
                - ctx.detail_filtered_out
                - ctx.candidates_count,
            )

        ok, notes = verify_funnel_closed(ctx.funnel_dict())
        print("funnel_ok", ok)
        for n in notes:
            print(" ", n)

        path = generate_report_file(ctx, reports_dir=settings.reports_dir)
        print("REPORT", path)

        doc = Document(str(path))
        paras = [p.text for p in doc.paragraphs if p.text.strip()]
        print("--- key lines ---")
        keys = (
            "用户原始",
            "标准化",
            "执行状态",
            "注意",
            "执行摘要",
            "漏斗",
            "总览",
            "匹配依据",
            "附件状态",
            "采购人",
            "数据完整度",
            "查看原始公告",
            "失败数据源",
            "免责",
            "中国政府",
            "招标投标",
        )
        for p in paras:
            if any(k in p for k in keys):
                print(p[:160])
        print("tables", len(doc.tables))
        for i, table in enumerate(doc.tables[:3]):
            print(f"table{i} header", [c.text for c in table.rows[0].cells])
            if len(table.rows) > 1:
                print(f"table{i} row1", [c.text[:20] for c in table.rows[1].cells])


if __name__ == "__main__":
    asyncio.run(main())

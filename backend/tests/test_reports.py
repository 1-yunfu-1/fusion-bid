"""Word 报告：结构、字段、统计闭合、超链接、文件可打开."""

from __future__ import annotations

from datetime import datetime
from pathlib import Path
from zipfile import ZipFile
from zoneinfo import ZoneInfo

import pytest
from docx import Document
from lxml import etree

from app.reports.fields import (
    attachment_status,
    enrich_report_item,
    extract_fields,
    format_cn_date,
    source_display_name,
    verify_funnel_closed,
)
from app.reports.filenames import build_report_filename, sanitize_query_for_filename
from app.reports.word_report import ReportContext, SourceRunStat, generate_report_file

TZ = ZoneInfo("Asia/Shanghai")
W_NS = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
R_NS = {"r": "http://schemas.openxmlformats.org/officeDocument/2006/relationships"}


def test_sanitize_filename_strips_illegal_and_traversal():
    q = '最近1个月../服务器<>:"/\\|?*招标'
    s = sanitize_query_for_filename(q)
    assert ".." not in s
    assert "<" not in s
    assert "/" not in s
    assert "\\" not in s
    assert "服务器" in s or "招标" in s


def test_build_report_filename_unique(tmp_path: Path):
    when = datetime(2026, 4, 7, 14, 24, tzinfo=TZ)
    p1 = build_report_filename(
        "最近3个月的上海区域内的充电桩招标信息都有哪些", when=when, reports_dir=tmp_path
    )
    assert p1.name.endswith("_202604071424.docx")
    assert "充电桩" in p1.name
    p1.write_bytes(b"pk")
    p2 = build_report_filename(
        "最近3个月的上海区域内的充电桩招标信息都有哪些", when=when, reports_dir=tmp_path
    )
    assert p2 != p1


def test_format_cn_date_and_source_names():
    assert format_cn_date("2026-07-16") == "2026年7月16日"
    assert format_cn_date(datetime(2026, 7, 16, tzinfo=TZ)) == "2026年7月16日"
    assert source_display_name("ccgp") == "中国政府采购网"
    assert source_display_name("cebpub") == "中国招标投标公共服务平台"


def test_extract_fields_no_fabrication():
    text = """
    项目名称：某某服务器采购项目
    采购人：北京市某局
    项目编号：BJ-2026-001
    预算金额：120万元
    投标截止时间：2026-07-20 09:00
    项目所在地：北京市海淀区
    采购内容：机架式服务器及配套
    资格要求：具备信息系统集成资质
    """
    f = extract_fields(
        title="北京市某局服务器采购公开招标公告",
        clean_content=text,
        publish_time="2026-07-10",
    )
    assert f["purchaser"] == "北京市某局"
    assert "BJ-2026-001" in f["project_code"]
    assert "120" in f["budget"]
    assert "北京市" in f["region"]
    assert f["publish_time_cn"] == "2026年7月10日"
    empty = extract_fields(title="仅有标题", clean_content="")
    assert empty["purchaser"] == "原文未明确说明"
    assert empty["budget"] == "原文未明确说明"


def test_attachment_status_variants():
    assert attachment_status(links=["http://a.pdf"])["status"] == "found"
    assert attachment_status(links=[], detail_fetched=False)["label"] == "详情未获取"
    assert attachment_status(links=[], extract_failed=True)["label"] == "提取失败"
    assert attachment_status(links=[], requires_login=True)["label"] == "登录后可见"
    assert attachment_status(links=[], detail_fetched=True)["label"] == "未发现公开附件"


def test_funnel_closed_ok_and_fail():
    ok, notes = verify_funnel_closed(
        {
            "raw_result_count": 10,
            "list_filtered_out": 3,
            "detail_cap_skipped": 1,
            "detail_failed": 1,
            "detail_filtered_out": 1,
            "candidates_count": 4,
            "cross_source_merge_count": 1,
            "primary_count": 3,
            "incremental_count": 2,
            "update_count": 1,
            "skipped_already_delivered": 0,
            "report_item_count": 3,
            "db_merge_count": 0,
        }
    )
    assert ok is True
    bad, _ = verify_funnel_closed(
        {
            "raw_result_count": 80,
            "list_filtered_out": 48,
            "detail_cap_skipped": 0,
            "detail_failed": 0,
            "detail_filtered_out": 0,
            "candidates_count": 8,
            "cross_source_merge_count": 0,
            "primary_count": 8,
            "incremental_count": 8,
            "update_count": 0,
            "skipped_already_delivered": 0,
            "report_item_count": 8,
        }
    )
    assert bad is False  # 80 != 48+8


def _sample_ctx(**kwargs) -> ReportContext:
    base = dict(
        original_query="最近1个月北京市服务器招标信息",
        generated_at=datetime(2026, 7, 17, 12, 0, tzinfo=TZ),
        execute_type="立即执行",
        data_mode="实时数据",
        execution_status="partial",
        keywords=["服务器"],
        regions=["北京市"],
        start_date="2026-06-17",
        end_date="2026-07-17",
        sources=["cebpub"],
        sources_requested=["ccgp", "cebpub"],
        sources_succeeded=["cebpub"],
        sources_failed={"ccgp": "连接超时或站点无响应"},
        source_stats=[
            SourceRunStat(
                source_name="ccgp",
                display_name="中国政府采购网",
                status="failed",
                message="连接超时或站点无响应",
                raw_count=0,
            ),
            SourceRunStat(
                source_name="cebpub",
                display_name="中国招标投标公共服务平台",
                status="success",
                message="检索完成",
                raw_count=10,
                list_kept=6,
                detail_success=5,
                final_contributed=2,
            ),
        ],
        raw_result_count=10,
        list_filtered_out=3,
        detail_cap_skipped=1,
        detail_failed=1,
        detail_success_count=5,
        detail_filtered_out=1,
        candidates_count=4,
        cross_source_merge_count=1,
        primary_count=3,
        filtered_out_count=4,
        duplicate_count=1,
        final_count=3,
        incremental_count=2,
        update_count=0,
        skipped_already_delivered=1,
        items=[
            {
                "title": "北京市海淀区教育系统服务器采购公开招标公告",
                "publish_time": "2026-07-01T00:00:00+08:00",
                "source_name": "cebpub",
                "source_url": "https://www.cebpubservice.com/example/1",
                "region": "北京市",
                "project_code": "BJ-HD-2026-88",
                "summary": "采购人：海淀区某中心\n预算金额：80万元",
                "clean_content": (
                    "采购人：海淀区某中心\n"
                    "项目编号：BJ-HD-2026-88\n"
                    "预算金额：80万元\n"
                    "投标截止时间：2026年7月20日\n"
                    "项目所在地：北京市海淀区\n"
                    "采购内容：机架式服务器30台\n"
                    "资格要求：注册资金不少于500万\n"
                    "代理机构：某某招标代理有限公司\n"
                ),
                "attachment_links": ["https://www.cebpubservice.com/a.pdf"],
                "change_label": "新增",
                "is_new": True,
                "is_update": False,
                "detail_fetched": True,
                "requires_login": False,
                "related_urls": [
                    "https://www.cebpubservice.com/example/1",
                    "https://www.ccgp.gov.cn/example/related",
                ],
            },
            {
                "title": "甘肃省某单位办公设备采购公告含服务器",
                "publish_time": "2026-07-05",
                "source_name": "cebpub",
                "source_url": "https://www.cebpubservice.com/example/2",
                "region": None,
                "clean_content": "采购内容：服务器及相关软件，交货地点：兰州市",
                "attachment_links": [],
                "change_label": "新增",
                "is_new": True,
                "detail_fetched": True,
                "requires_login": False,
            },
        ],
        crawl_time="2026-07-17T12:00:00+08:00",
        warnings=["存在失败或跳过的数据源，报告覆盖范围不完整。"],
    )
    base.update(kwargs)
    return ReportContext(**base)


def test_generate_word_report_structure_and_fields(tmp_path: Path):
    ctx = _sample_ctx()
    path = generate_report_file(ctx, reports_dir=tmp_path)
    assert path.is_file()
    assert path.suffix == ".docx"
    # 文件可被 python-docx 打开
    doc = Document(str(path))
    text = "\n".join(p.text for p in doc.paragraphs)
    assert "用户原始问题" in text
    assert "标准化查询条件" in text or "关键词" in text
    assert "北京市" in text
    assert "执行摘要" in text or "数据处理漏斗" in text
    assert "中国政府采购网" in text
    assert "中国招标投标公共服务平台" in text
    assert "部分成功" in text or "注意" in text
    assert "2026年7月1日" in text or "2026年7月" in text
    assert "海淀区某中心" in text  # 采购人抽取
    assert "匹配依据" in text
    assert "数据完整度" in text
    assert "查看原始公告" in text
    assert "https://www.cebpubservice.com/example/1" not in text  # 不直接铺满长 URL
    assert "未发现公开附件" in text or "发现" in text or "附件状态" in text
    assert "数据处理说明" in text or "免责声明" in text
    assert "连接超时" in text or "失败数据源" in text
    # 页眉
    assert doc.sections[0].header.paragraphs
    assert any("智标" in (p.text or "") or "FusionBid" in (p.text or "") for p in doc.sections[0].header.paragraphs)


def test_report_hyperlinks_clickable(tmp_path: Path):
    path = generate_report_file(_sample_ctx(), reports_dir=tmp_path)
    # 校验 document.xml.rels 含 external hyperlink
    with ZipFile(path, "r") as zf:
        rels = zf.read("word/_rels/document.xml.rels").decode("utf-8")
        assert "hyperlink" in rels.lower()
        assert "cebpubservice.com/example/1" in rels
        # 正文含超链接文字
        xml = zf.read("word/document.xml")
    root = etree.fromstring(xml)
    texts = [t.text or "" for t in root.xpath(".//w:t", namespaces=W_NS)]
    joined = "".join(texts)
    assert "查看原始公告" in joined
    # hyperlink 元素存在
    links = root.xpath(".//w:hyperlink", namespaces=W_NS)
    assert len(links) >= 1


def test_report_stats_closed_in_document(tmp_path: Path):
    ctx = _sample_ctx()
    ok, notes = verify_funnel_closed(ctx.funnel_dict())
    assert ok, notes
    path = generate_report_file(ctx, reports_dir=tmp_path)
    doc = Document(str(path))
    text = "\n".join(p.text for p in doc.paragraphs)
    # 表格也算
    for table in doc.tables:
        for row in table.rows:
            text += "\n" + " | ".join(c.text for c in row.cells)
    assert "原始列表结果" in text
    assert "本报告条目" in text
    assert "闭合校验通过" in text or "10 = 3+1+1+1+4" in text or "闭合" in text


def test_enrich_match_basis_region_warning():
    item = enrich_report_item(
        {
            "title": "甘肃省兰州服务器采购公告",
            "clean_content": "采购服务器一批",
            "publish_time": "2026-07-01",
            "source_name": "cebpub",
            "source_url": "https://example.com/x",
            "attachment_links": [],
            "detail_fetched": True,
        },
        keywords=["服务器"],
        regions=["北京市"],
        start_date="2026-06-01",
        end_date="2026-07-17",
    )
    assert "甘肃" in item["match_basis"]["region"] or "人工复核" in item["match_basis"]["region"]
    assert item["attachment"]["label"] == "未发现公开附件"
    assert "percent" in item["completeness"]


@pytest.mark.asyncio
async def test_reports_list_api(client, tmp_path, monkeypatch):
    from app.core.config import get_settings

    settings = get_settings()
    monkeypatch.setattr(settings, "reports_dir", tmp_path)
    path = generate_report_file(
        ReportContext(
            original_query="测试报告查询",
            items=[],
            raw_result_count=0,
            candidates_count=0,
            primary_count=0,
            final_count=0,
        ),
        reports_dir=tmp_path,
    )
    resp = await client.get("/api/reports")
    assert resp.status_code == 200
    data = resp.json()
    assert data["total"] >= 1
    names = [i["filename"] for i in data["items"]]
    assert path.name in names

    dl = await client.get(f"/api/reports/download/{path.name}")
    assert dl.status_code == 200


@pytest.mark.asyncio
async def test_download_rejects_path_traversal(client):
    resp = await client.get("/api/reports/download/..%2F..%2Fetc%2Fpasswd.docx")
    assert resp.status_code in (400, 404)

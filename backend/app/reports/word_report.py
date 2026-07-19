"""Word 招投标报告生成（结构化、可审计、统计闭合）."""

from __future__ import annotations

from dataclasses import dataclass, field
from datetime import datetime
from pathlib import Path
from typing import Any
from zoneinfo import ZoneInfo

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor

from app.reports.fields import (
    _MISSING,
    enrich_report_item,
    format_cn_date,
    format_cn_datetime,
    source_display_name,
    verify_funnel_closed,
)
from app.reports.filenames import build_report_filename

TZ = ZoneInfo("Asia/Shanghai")


@dataclass
class SourceRunStat:
    """单数据源执行情况."""

    source_name: str
    display_name: str = ""
    status: str = "unknown"  # success | failed | skipped
    message: str = ""
    raw_count: int = 0
    list_kept: int = 0
    detail_success: int = 0
    detail_metadata_only: int = 0
    final_contributed: int = 0  # 进入报告的条目数（按主来源）


@dataclass
class ReportContext:
    system_name: str = "FusionBid 智标聚合助手"
    original_query: str = ""
    generated_at: datetime = field(default_factory=lambda: datetime.now(TZ))
    execute_type: str = "立即执行"
    data_mode: str = "实时数据"
    execution_status: str = "success"  # success | partial | failed
    keywords: list[str] = field(default_factory=list)
    regions: list[str] = field(default_factory=list)
    start_date: str | None = None
    end_date: str | None = None
    schedule_desc: str = "无（立即执行）"
    # 兼容旧字段：成功源标识列表
    sources: list[str] = field(default_factory=list)
    sources_requested: list[str] = field(default_factory=list)
    sources_succeeded: list[str] = field(default_factory=list)
    sources_failed: dict[str, str] = field(default_factory=dict)
    source_stats: list[SourceRunStat] = field(default_factory=list)
    # 闭合漏斗
    raw_result_count: int = 0
    list_filtered_out: int = 0
    detail_cap: int = 30
    detail_cap_skipped: int = 0
    detail_failed: int = 0
    detail_success_count: int = 0
    detail_metadata_only_count: int = 0
    detail_status_failed_count: int = 0
    detail_human_verification_count: int = 0
    detail_filtered_out: int = 0
    candidates_count: int = 0
    cross_source_merge_count: int = 0
    primary_count: int = 0
    db_merge_count: int = 0
    filtered_out_count: int = 0  # 兼容：list+detail 过滤合计
    duplicate_count: int = 0
    final_count: int = 0  # 去重后主记录
    incremental_count: int = 0
    update_count: int = 0
    skipped_already_delivered: int = 0
    report_scope: str = "incremental"
    truncated: bool = False
    deduplicate: bool = True
    analysis: dict[str, Any] = field(default_factory=dict)
    # 结果
    items: list[dict[str, Any]] = field(default_factory=list)
    crawl_time: str | None = None
    extra_notes: list[str] = field(default_factory=list)
    warnings: list[str] = field(default_factory=list)

    def funnel_dict(self) -> dict[str, int]:
        return {
            "raw_result_count": self.raw_result_count,
            "list_filtered_out": self.list_filtered_out,
            "detail_cap_skipped": self.detail_cap_skipped,
            "detail_failed": self.detail_failed,
            "detail_filtered_out": self.detail_filtered_out,
            "candidates_count": self.candidates_count,
            "cross_source_merge_count": self.cross_source_merge_count,
            "primary_count": self.primary_count or self.final_count,
            "incremental_count": self.incremental_count,
            "update_count": self.update_count,
            "skipped_already_delivered": self.skipped_already_delivered,
            "report_item_count": len(self.items),
            "db_merge_count": self.db_merge_count,
        }


def _set_run_font(
    run,
    *,
    size: int = 11,
    bold: bool = False,
    name: str = "宋体",
    color: RGBColor | None = None,
) -> None:
    run.bold = bold
    run.font.size = Pt(size)
    run.font.name = name
    r = run._element
    rPr = r.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.insert(0, rFonts)
    rFonts.set(qn("w:eastAsia"), name)
    rFonts.set(qn("w:ascii"), name if name != "宋体" else "Times New Roman")
    if color is not None:
        run.font.color.rgb = color


def _add_heading(doc: Document, text: str, level: int = 1) -> None:
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        _set_run_font(run, size=16 if level == 1 else 13, bold=True, name="黑体")


def _add_para(
    doc: Document,
    text: str,
    *,
    bold: bool = False,
    size: int = 11,
    color: RGBColor | None = None,
    center: bool = False,
) -> Any:
    p = doc.add_paragraph()
    if center:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    _set_run_font(run, size=size, bold=bold, color=color)
    return p


def _add_label_value(doc: Document, label: str, value: str, *, size: int = 11) -> None:
    p = doc.add_paragraph()
    r1 = p.add_run(f"{label}")
    _set_run_font(r1, size=size, bold=True)
    r2 = p.add_run(str(value) if value is not None else _MISSING)
    _set_run_font(r2, size=size)


def _add_bullet(doc: Document, text: str, *, size: int = 10) -> None:
    """Use a real Word bullet paragraph for source-derived requirement lists."""
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Cm(1.27)  # 0.5 in
    p.paragraph_format.first_line_indent = Cm(-0.635)  # hanging 0.25 in
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.line_spacing = 1.167
    run = p.add_run(text)
    _set_run_font(run, size=size)


def _add_hyperlink(paragraph, url: str, text: str = "查看原始公告") -> None:
    part = paragraph.part
    r_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)
    new_run = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "0563C1")
    rPr.append(color)
    u = OxmlElement("w:u")
    u.set(qn("w:val"), "single")
    rPr.append(u)
    sz = OxmlElement("w:sz")
    sz.set(qn("w:val"), "22")
    rPr.append(sz)
    rFonts = OxmlElement("w:rFonts")
    rFonts.set(qn("w:eastAsia"), "宋体")
    rPr.append(rFonts)
    new_run.append(rPr)
    t = OxmlElement("w:t")
    t.text = text
    new_run.append(t)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)


def _set_cell_shading(cell, fill: str) -> None:
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    shd.set(qn("w:val"), "clear")
    tcPr.append(shd)


def _set_cell_text(cell, text: str, *, bold: bool = False, size: int = 9) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    run = p.add_run(str(text) if text is not None else "")
    _set_run_font(run, size=size, bold=bold)


def _set_table_geometry(table, widths: list[int]) -> None:
    """Apply the standard_business_brief fixed 9360-DXA table token map."""
    table.autofit = False
    tbl_pr = table._tbl.tblPr

    def child(tag: str):
        value = tbl_pr.find(qn(tag))
        if value is None:
            value = OxmlElement(tag)
            tbl_pr.append(value)
        return value

    table_width = child("w:tblW")
    table_width.set(qn("w:w"), "9360")
    table_width.set(qn("w:type"), "dxa")
    table_indent = child("w:tblInd")
    table_indent.set(qn("w:w"), "120")
    table_indent.set(qn("w:type"), "dxa")
    layout = child("w:tblLayout")
    layout.set(qn("w:type"), "fixed")

    grid_cols = table._tbl.tblGrid.findall(qn("w:gridCol"))
    for index, width in enumerate(widths):
        if index < len(grid_cols):
            grid_cols[index].set(qn("w:w"), str(width))
    for row in table.rows:
        for index, cell in enumerate(row.cells):
            width = widths[min(index, len(widths) - 1)]
            cell.width = Inches(width / 1440)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            tc_mar = tc_pr.find(qn("w:tcMar"))
            if tc_mar is None:
                tc_mar = OxmlElement("w:tcMar")
                tc_pr.append(tc_mar)
            for side, value in (("top", 80), ("bottom", 80), ("start", 120), ("end", 120)):
                margin = tc_mar.find(qn(f"w:{side}"))
                if margin is None:
                    margin = OxmlElement(f"w:{side}")
                    tc_mar.append(margin)
                margin.set(qn("w:w"), str(value))
                margin.set(qn("w:type"), "dxa")


def _add_table(doc: Document, headers: list[str], rows: list[list[str]], *, col_widths_cm: list[float] | None = None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        _set_cell_text(cell, h, bold=True, size=9)
        try:
            _set_cell_shading(cell, "D9E8F5")
        except Exception:  # noqa: BLE001
            pass
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            _set_cell_text(table.rows[ri + 1].cells[ci], val, size=9)
    raw_widths = col_widths_cm or [1.0] * len(headers)
    total = sum(raw_widths) or float(len(headers))
    widths = [max(1, round(9360 * width / total)) for width in raw_widths]
    # Keep the table grid exactly at 9360 DXA despite rounding individual columns.
    widths[-1] += 9360 - sum(widths)
    _set_table_geometry(table, widths)
    return table


def _status_label(status: str) -> str:
    return {
        "success": "成功（全部请求数据源已完成）",
        "partial": "部分成功（存在失败或跳过的数据源）",
        "failed": "失败",
        "running": "执行中",
    }.get(status, status or "未知")


def _period_text(ctx: ReportContext) -> str:
    if ctx.start_date or ctx.end_date:
        a = format_cn_date(ctx.start_date) if ctx.start_date else "不限"
        b = format_cn_date(ctx.end_date) if ctx.end_date else "不限"
        return f"{a} 至 {b}"
    return "未指定统计周期"


def _prepare_items(ctx: ReportContext) -> list[dict[str, Any]]:
    analysis_by_id = {
        str(project.get("announcement_id")): project
        for project in (ctx.analysis or {}).get("projects", [])
        if isinstance(project, dict) and project.get("announcement_id")
    }
    prepared = [
        enrich_report_item(
            it,
            keywords=ctx.keywords,
            regions=ctx.regions,
            start_date=ctx.start_date,
            end_date=ctx.end_date,
        )
        for it in ctx.items
    ]
    for item in prepared:
        item["opportunity_analysis"] = analysis_by_id.get(
            str(item.get("announcement_id")), {}
        )
    return prepared


def _add_header_footer(doc: Document, ctx: ReportContext) -> None:
    section = doc.sections[0]
    # standard_business_brief：Letter + 1 in margins + 0.492 in header/footer.
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(2.54)
    section.right_margin = Cm(2.54)
    section.header_distance = Cm(1.25)
    section.footer_distance = Cm(1.25)

    header = section.header
    header.is_linked_to_previous = False
    hp = header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    hr = hp.add_run(f"{ctx.system_name} · 招投标信息聚合报告")
    _set_run_font(hr, size=9, name="宋体", color=RGBColor(0x66, 0x66, 0x66))

    footer = section.footer
    footer.is_linked_to_previous = False
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER

    def _add_page_field(paragraph, instr: str) -> None:
        run = paragraph.add_run()
        fld_begin = OxmlElement("w:fldChar")
        fld_begin.set(qn("w:fldCharType"), "begin")
        instr_text = OxmlElement("w:instrText")
        instr_text.set(qn("xml:space"), "preserve")
        instr_text.text = instr
        fld_sep = OxmlElement("w:fldChar")
        fld_sep.set(qn("w:fldCharType"), "separate")
        fld_end = OxmlElement("w:fldChar")
        fld_end.set(qn("w:fldCharType"), "end")
        run._r.append(fld_begin)
        run._r.append(instr_text)
        run._r.append(fld_sep)
        run._r.append(fld_end)

    r0 = fp.add_run("第 ")
    _set_run_font(r0, size=9, color=RGBColor(0x66, 0x66, 0x66))
    _add_page_field(fp, " PAGE ")
    r1 = fp.add_run(" 页 / 共 ")
    _set_run_font(r1, size=9, color=RGBColor(0x66, 0x66, 0x66))
    _add_page_field(fp, " NUMPAGES ")
    r2 = fp.add_run(" 页")
    _set_run_font(r2, size=9, color=RGBColor(0x66, 0x66, 0x66))


def _page1_cover(doc: Document, ctx: ReportContext, items: list[dict]) -> None:
    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t.add_run(ctx.system_name)
    _set_run_font(r, size=22, bold=True, name="黑体")

    s = doc.add_paragraph()
    s.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = s.add_run("招投标信息聚合报告")
    _set_run_font(r2, size=18, bold=True, name="黑体")

    doc.add_paragraph()
    _add_label_value(doc, "用户原始问题：", ctx.original_query or _MISSING, size=12)
    _add_label_value(
        doc,
        "标准化查询条件：",
        (
            f"关键词：{'、'.join(ctx.keywords) if ctx.keywords else '（未指定）'}；"
            f"区域：{'、'.join(ctx.regions) if ctx.regions else '（未指定）'}；"
            f"周期：{_period_text(ctx)}；"
            f"执行方式：{ctx.execute_type}；"
            f"频率：{ctx.schedule_desc}"
        ),
    )
    _add_label_value(doc, "统计周期：", _period_text(ctx))
    _add_label_value(doc, "有效结果数量（本报告条目）：", str(len(items)))
    succ = len(ctx.sources_succeeded or ctx.sources)
    req = len(ctx.sources_requested) or succ
    _add_label_value(
        doc,
        "成功数据源数量：",
        f"{succ} / {req}" if req else str(succ),
    )
    _add_label_value(doc, "本次新增数量：", str(ctx.incremental_count))
    _add_label_value(doc, "本次内容更新数量：", str(ctx.update_count))
    _add_label_value(
        doc,
        "报告范围：",
        (
            "本轮未去重快照（已达安全上限，结果截断）"
            if ctx.truncated
            else "本轮原始完整快照（不去重）"
        )
        if ctx.report_scope == "snapshot"
        else "增量交付（新增与变化）",
    )
    _add_label_value(doc, "数据模式：", ctx.data_mode)
    _add_label_value(
        doc,
        "报告生成时间：",
        format_cn_datetime(ctx.generated_at),
    )

    st = ctx.execution_status or "success"
    p = doc.add_paragraph()
    r1 = p.add_run("执行状态：")
    _set_run_font(r1, size=12, bold=True)
    color = {
        "success": RGBColor(0x1E, 0x84, 0x4E),
        "partial": RGBColor(0xB7, 0x6E, 0x00),
        "failed": RGBColor(0xC0, 0x39, 0x2B),
    }.get(st, RGBColor(0x33, 0x33, 0x33))
    r2 = p.add_run(_status_label(st))
    _set_run_font(r2, size=12, bold=True, color=color)

    if st == "partial":
        box = doc.add_paragraph()
        run = box.add_run(
            "【注意】本次为部分成功：可能存在数据源失败、详情未验证、"
            "待人工验证或快照截断。请结合「数据源执行情况」和每条「详情状态」阅读，"
            "勿将结果理解为全网完整覆盖。"
        )
        _set_run_font(run, size=11, bold=True, color=RGBColor(0xB7, 0x6E, 0x00))

    if ctx.warnings:
        _add_heading(doc, "警告与提示", 2)
        for w in ctx.warnings:
            _add_para(doc, f"· {w}", color=RGBColor(0xB7, 0x6E, 0x00))

    if ctx.data_mode != "实时数据":
        _add_para(
            doc,
            "【注意】本报告包含演示/测试数据，请勿当作全部 live 结果。",
            bold=True,
            color=RGBColor(0xC0, 0x39, 0x2B),
        )


def _page2_summary(doc: Document, ctx: ReportContext, items: list[dict]) -> None:
    doc.add_page_break()
    _add_heading(doc, "一、执行摘要", 1)
    src_ok = "、".join(source_display_name(s) for s in (ctx.sources_succeeded or ctx.sources)) or "无"
    src_fail = "、".join(
        source_display_name(s) for s in (ctx.sources_failed or {})
    ) or "无"
    report_item_summary = (
        f"本报告保留来源记录 {len(items)} 条（未去重，可能重复项已标注）"
        if ctx.report_scope == "snapshot"
        else (
            f"本报告收录有效结果 {len(items)} 条"
            f"（新增 {ctx.incremental_count}、更新 {ctx.update_count}）"
        )
    )
    _add_para(
        doc,
        (
            f"围绕「{ctx.original_query or '（空查询）'}」，"
            f"在统计周期 {_period_text(ctx)} 内，"
            f"共请求 {len(ctx.sources_requested) or '—'} 个数据源，"
            f"成功 {len(ctx.sources_succeeded or ctx.sources)} 个"
            f"（{src_ok}），"
            f"失败/跳过 {len(ctx.sources_failed or {})} 个（{src_fail}）。"
            f"原始列表 {ctx.raw_result_count} 条，"
            f"经过滤与去重后主记录 {ctx.primary_count or ctx.final_count} 条，"
            f"{report_item_summary}。"
        ),
    )

    _add_heading(doc, "二、数据源执行情况", 1)
    rows: list[list[str]] = []
    if ctx.source_stats:
        for ss in ctx.source_stats:
            rows.append(
                [
                    ss.display_name or source_display_name(ss.source_name),
                    {"success": "成功", "failed": "失败", "skipped": "跳过"}.get(
                        ss.status, ss.status
                    ),
                    str(ss.raw_count),
                    str(ss.detail_success),
                    str(ss.detail_metadata_only),
                    str(ss.final_contributed),
                    (ss.message or "—")[:40],
                ]
            )
    else:
        # 回退：从 succeeded/failed 拼表
        all_names = list(
            dict.fromkeys(
                (ctx.sources_requested or [])
                + (ctx.sources_succeeded or ctx.sources or [])
                + list((ctx.sources_failed or {}).keys())
            )
        )
        for name in all_names:
            if name in (ctx.sources_failed or {}):
                st = "失败"
                msg = (ctx.sources_failed or {}).get(name, "")[:40]
            elif name in (ctx.sources_succeeded or ctx.sources or []):
                st = "成功"
                msg = "已完成检索"
            else:
                st = "未执行"
                msg = "—"
            contrib = sum(1 for it in items if it.get("source_name") == name)
            rows.append([source_display_name(name), st, "—", "—", "—", str(contrib), msg])
    if not rows:
        rows.append(["—", "—", "0", "0", "0", "0", "无数据源记录"])
    _add_table(
        doc,
        ["数据源", "状态", "列表条数", "完整详情", "仅元数据", "报告贡献", "说明"],
        rows,
        col_widths_cm=[3.2, 1.3, 1.5, 1.5, 1.5, 1.5, 3.7],
    )
    _add_para(
        doc,
        "说明：「完整详情」仅计入已通过当前公告校验的详情；「仅元数据」不会被当作公告正文。"
        "「报告贡献」指本报告条目中主来源为该数据源的数量。",
        size=9,
        color=RGBColor(0x66, 0x66, 0x66),
    )

    _add_heading(doc, "三、数据处理漏斗", 1)
    fd = ctx.funnel_dict()
    if ctx.report_scope == "snapshot":
        ok, notes = True, [
            (
                "未去重快照已达每源 500 条安全上限，已标记截断，不声称为完整报告。"
                if ctx.truncated
                else "未去重完整快照保留本轮全部合格来源记录；本次不读取也不写入 DeliveryHistory。"
            )
        ]
    else:
        ok, notes = verify_funnel_closed(fd)
    funnel_rows = [
        ["原始列表结果", str(ctx.raw_result_count), "各成功数据源列表页合计"],
        ["列表阶段过滤排除", str(ctx.list_filtered_out), "关键词/区域/时间/相关性"],
        [
            "详情上限未抓取",
            str(ctx.detail_cap_skipped),
            f"每源最多采集 {ctx.detail_cap} 条详情",
        ],
        ["详情抓取失败", str(ctx.detail_failed), "打开详情页异常"],
        ["仅列表元数据", str(ctx.detail_metadata_only_count), "未验证详情，不作为公告正文"],
        ["详情内容失败", str(ctx.detail_status_failed_count), "已保留来源记录与失败状态"],
        ["待人工安全验证", str(ctx.detail_human_verification_count), "系统未绕过验证机制"],
        ["详情阶段过滤排除", str(ctx.detail_filtered_out), "详情二次条件不匹配"],
        ["进入去重候选", str(ctx.candidates_count), "待跨源去重"],
        ["跨源合并减少", str(ctx.cross_source_merge_count), "同项目多来源合并"],
        ["去重后主记录", str(ctx.primary_count or ctx.final_count), "本批唯一主记录"],
        ["库内已存在合并", str(ctx.db_merge_count), "与历史库记录合并"],
        ["历史已交付跳过", str(ctx.skipped_already_delivered), "增量：已推送且无变化"],
        ["本次新增", str(ctx.incremental_count), "首次交付"],
        ["本次内容更新", str(ctx.update_count), "已推送但内容变化"],
        [
            "本报告条目",
            str(len(items)),
            "本轮全部来源记录（未去重）"
            if ctx.report_scope == "snapshot"
            else "新增 + 更新",
        ],
    ]
    _add_table(
        doc,
        ["阶段", "数量", "说明"],
        funnel_rows,
        col_widths_cm=[4.5, 2.0, 9.0],
    )
    for n in notes:
        _add_para(
            doc,
            f"{'✓' if ok else '!'} {n}",
            size=9,
            color=RGBColor(0x1E, 0x84, 0x4E) if ok else RGBColor(0xC0, 0x39, 0x2B),
        )

    _add_heading(doc, "四、投标决策分析", 1)
    analysis = ctx.analysis or {}
    _add_para(
        doc,
        str(analysis.get("portfolio_summary") or "本轮未生成组合研判。"),
        bold=True,
    )
    _add_para(
        doc,
        "分析方式：规则依据公告原文字段生成；如启用 LLM，仅展示通过字段证据校验的补充建议，不构成中标预测。",
        size=9,
        color=RGBColor(0x66, 0x66, 0x66),
    )
    projects = [
        p
        for p in analysis.get("projects", [])
        if isinstance(p, dict) and p.get("is_opportunity", True)
    ]
    if projects:
        analysis_rows = []
        for project in projects[:8]:
            analysis_rows.append(
                [
                    str(project.get("title") or _MISSING)[:36],
                    str(project.get("decision") or "信息不足"),
                    str(project.get("deadline_urgency") or "未知"),
                    "；".join(project.get("gaps") or ["无明显缺口"])[:60],
                    "；".join(project.get("recommended_actions") or [])[:80],
                ]
            )
        _add_table(
            doc,
            ["项目", "参与建议", "时效", "主要信息缺口", "建议下一步"],
            analysis_rows,
            col_widths_cm=[4.0, 1.5, 1.6, 4.0, 5.1],
        )
    else:
        _add_para(doc, "暂无可分析项目；可先补齐官方详情后重新生成未去重完整报告。")

    _add_heading(doc, "五、重点项目 Top 3", 1)
    opportunity_items = [
        it
        for it in items
        if (it.get("fields") or {}).get("lifecycle_stage") == "机会公告"
    ]
    if not opportunity_items:
        _add_para(doc, "本次无有效结果条目。")
    else:
        # 优先级由证据受限的规则研判给出，再按完整度排序。
        priority_rank = {"高": 4, "中": 3, "待核验": 2, "低": 1}
        ranked = sorted(
            opportunity_items,
            key=lambda x: (
                priority_rank.get(
                    (x.get("opportunity_analysis") or {}).get("priority"), 0
                ),
                x.get("completeness", {}).get("percent") or 0,
            ),
            reverse=True,
        )[:3]
        for i, it in enumerate(ranked, 1):
            f = it.get("fields") or {}
            _add_para(
                doc,
                f"{i}. {it.get('short_title') or it.get('title') or _MISSING}",
                bold=True,
            )
            _add_para(
                doc,
                (
                    f"地区：{f.get('region', _MISSING)}；"
                    f"{f.get('purchaser_source_label') or '采购人/招标人'}：{f.get('purchaser', _MISSING)}；"
                    f"截止：{f.get('deadline', _MISSING)}；"
                    f"优先级：{(it.get('opportunity_analysis') or {}).get('priority', '待核验')}；"
                    f"完整度：{it.get('completeness', {}).get('label', '—')}；"
                    f"来源：{it.get('source_display', '—')}"
                ),
                size=10,
            )

    _add_heading(doc, "六、数据质量提示", 1)
    tips = [
        "所有字段均来自公开页面抽取；无法确认时标注「原文未明确说明」或「本次未成功提取」，系统不编造。",
        "区域匹配依据见各项目「匹配依据」；标题含外地地名时请人工核对是否误召回。",
        "附件状态区分：未发现公开附件 / 提取失败 / 登录后可见 / 详情未获取，不再统一写「无」。",
        "部分成功时，失败数据源的公告不会出现在本报告中。",
    ]
    tips.extend(ctx.extra_notes)
    for t in tips:
        _add_para(doc, f"· {t}")


def _page3_overview(doc: Document, ctx: ReportContext, items: list[dict]) -> None:
    doc.add_page_break()
    _add_heading(doc, "七、项目总览", 1)
    if not items:
        _add_para(doc, "无项目可展示。")
        return
    opportunity_items = [
        it for it in items if (it.get("fields") or {}).get("lifecycle_stage") == "机会公告"
    ]
    lifecycle_items = [it for it in items if it not in opportunity_items]

    _add_heading(doc, "七（一）、可参与机会", 2)
    rows = []
    for i, it in enumerate(opportunity_items, 1):
        f = it.get("fields") or {}
        cached_full_reused = bool(it.get("cached_full_reused"))
        rows.append(
            [
                str(i),
                it.get("short_title") or _MISSING,
                f.get("region") or _MISSING,
                f.get("purchaser") or _MISSING,
                (it.get("opportunity_analysis") or {}).get("priority") or "待核验",
                (
                    "历史完整正文（本轮失败）"
                    if cached_full_reused
                    else {
                        "full": "已核验详情",
                        "metadata_only": "仅元数据",
                        "failed": "详情失败",
                        "needs_human_verification": "待人工验证",
                    }.get(it.get("detail_status"), "状态未知")
                ),
                f.get("procurement_method") or _MISSING,
                it.get("publish_time_cn") or _MISSING,
                f.get("deadline") or _MISSING,
                f.get("budget") or _MISSING,
                it.get("source_display") or source_display_name(it.get("source_name")),
            ]
        )
    _add_table(
        doc,
        [
            "序号",
            "项目简称",
            "项目地区",
            "采购人/招标人",
            "优先级",
            "详情状态",
            "采购方式",
            "发布时间",
            "截止时间",
            "预算金额",
            "来源",
        ],
        rows,
        col_widths_cm=[0.8, 2.4, 1.4, 1.8, 1.1, 1.5, 1.2, 1.4, 1.4, 1.3, 1.4],
    )
    if not rows:
        _add_para(doc, "本次报告没有可参与机会。")

    _add_heading(doc, "七（二）、项目生命周期情报", 2)
    lifecycle_rows = []
    for i, it in enumerate(lifecycle_items, 1):
        f = it.get("fields") or {}
        stage = f.get("lifecycle_stage") or "待复核"
        key_fact = (
            f"中标人：{f.get('awardee') or _MISSING}；金额：{f.get('award_amount') or _MISSING}"
            if stage == "结果公告"
            else f.get("change_summary") or _MISSING
            if stage == "更正/澄清"
            else f.get("termination_reason") or _MISSING
        )
        lifecycle_rows.append(
            [
                str(i),
                it.get("short_title") or _MISSING,
                stage,
                f.get("project_code") or _MISSING,
                str(key_fact)[:100],
                it.get("publish_time_cn") or _MISSING,
                it.get("source_display") or source_display_name(it.get("source_name")),
            ]
        )
    if lifecycle_rows:
        _add_table(
            doc,
            ["序号", "项目简称", "生命周期", "项目编号", "关键情报", "发布时间", "来源"],
            lifecycle_rows,
            col_widths_cm=[0.8, 3.0, 1.8, 2.4, 5.0, 1.8, 2.0],
        )
    else:
        _add_para(doc, "本次报告没有结果、更正或终止类生命周期情报。")


def _detail_pages(doc: Document, ctx: ReportContext, items: list[dict]) -> None:
    doc.add_page_break()
    _add_heading(doc, "八、项目结构化详情", 1)
    if not items:
        _add_para(doc, "无详情可展示。")
        return

    for i, it in enumerate(items, 1):
        f = it.get("fields") or {}
        title = it.get("title") or _MISSING
        short = it.get("short_title") or title
        label = it.get("change_label") or ""
        head = f"{i}. {short}"
        if label:
            head = f"{head} 【{label}】"
        _add_heading(doc, head, 2)

        # 完整标题仅显示一次
        if title and title != short:
            _add_label_value(doc, "完整标题：", title)
        # 项目名称若与标题/简称重复则跳过
        pn = f.get("project_name_display") or ""
        if pn and pn not in (title, short):
            _add_label_value(doc, "项目名称：", pn)

        status_bits: list[str] = []
        if label:
            status_bits.append(label)
        if it.get("is_update"):
            status_bits.append("内容更新")
        elif it.get("is_new"):
            status_bits.append("新增")
        status_bits = list(dict.fromkeys(status_bits))
        _add_label_value(
            doc,
            "公告状态：",
            "、".join(status_bits) if status_bits else "有效收录",
        )
        if it.get("dedupe_status"):
            dedupe_text = str(it["dedupe_status"])
            if it.get("dedupe_hint"):
                dedupe_text = f"{dedupe_text}：{it['dedupe_hint']}"
            _add_label_value(doc, "收录方式：", dedupe_text)
        _add_label_value(doc, "项目编号：", f.get("project_code") or _MISSING)
        purchaser_label = f.get("purchaser_source_label")
        if not purchaser_label or purchaser_label == _MISSING:
            purchaser_label = "采购人/招标人"
        _add_label_value(doc, f"{purchaser_label}：", f.get("purchaser") or _MISSING)
        if (
            f.get("tenderer")
            and f.get("tenderer") not in {_MISSING, f.get("purchaser")}
        ):
            tenderer_label = f.get("tenderer_source_label") or "招标人"
            _add_label_value(doc, f"{tenderer_label}：", f.get("tenderer"))
        _add_label_value(doc, "招标/采购代理机构：", f.get("agency") or _MISSING)
        _add_label_value(doc, "交易平台：", f.get("transaction_platform") or _MISSING)
        _add_label_value(doc, "项目地区：", f.get("region") or _MISSING)
        lifecycle_stage = f.get("lifecycle_stage") or "待复核"
        _add_label_value(doc, "公告生命周期：", lifecycle_stage)
        _add_label_value(doc, "采购方式：", f.get("procurement_method") or _MISSING)
        _add_label_value(doc, "采购内容：", f.get("content") or _MISSING)
        if lifecycle_stage == "结果公告":
            _add_label_value(doc, "中标人/成交供应商：", f.get("awardee") or _MISSING)
            _add_label_value(doc, "中标/成交金额：", f.get("award_amount") or _MISSING)
        elif lifecycle_stage == "更正/澄清":
            _add_label_value(doc, "更正/澄清事项：", f.get("change_summary") or _MISSING)
        elif lifecycle_stage == "终止/废标":
            _add_label_value(doc, "终止/废标原因：", f.get("termination_reason") or _MISSING)
        else:
            _add_label_value(doc, "预算金额：", f.get("budget") or _MISSING)
            _add_label_value(doc, "招标文件售价：", f.get("document_price") or _MISSING)
        _add_label_value(doc, "资金来源：", f.get("funding_source") or _MISSING)
        _add_label_value(doc, "发布时间：", it.get("publish_time_cn") or _MISSING)
        _add_label_value(doc, "公告结束时间：", f.get("notice_end_time") or _MISSING)
        _add_label_value(
            doc,
            "招标文件获取时间：",
            (
                f"{f.get('document_acquisition_start')} 至 {f.get('document_acquisition_end')}"
                if f.get("document_acquisition_start") not in {None, "", _MISSING}
                and f.get("document_acquisition_end") not in {None, "", _MISSING}
                else f.get("document_acquisition_end") or f.get("document_acquisition_start") or _MISSING
            ),
        )
        _add_label_value(doc, "投标截止时间：", f.get("bid_deadline") or f.get("deadline") or _MISSING)
        _add_label_value(doc, "开标时间：", f.get("opening_time") or _MISSING)
        qualification_items = f.get("qualification_items") or []
        if lifecycle_stage != "机会公告":
            pass
        elif qualification_items:
            _add_label_value(doc, "资格要求（原文提取）：", "见以下条款")
            for requirement in qualification_items:
                _add_bullet(doc, str(requirement), size=10)
        else:
            _add_label_value(doc, "资格要求：", f.get("qualification") or _MISSING)
        _add_label_value(
            doc,
            "详情状态：",
            (
                "使用历史已核验完整正文（本轮采集未成功，不计为本轮完整详情）"
                if it.get("cached_full_reused")
                else {
                    "full": "已核验公告详情",
                    "metadata_only": "仅列表元数据（不作为公告正文）",
                    "failed": "详情抓取失败",
                    "needs_human_verification": "详情页需人工完成安全验证（系统未绕过）",
                }.get(it.get("detail_status"), "状态未知")
            ),
        )
        if it.get("cached_full_reused") or it.get("detail_status") != "full":
            source_metadata = it.get("source_metadata") or {}
            last_attempt = (
                source_metadata.get("last_skip")
                or source_metadata.get("last_attempt")
                or source_metadata
            )
            attempt_reason = last_attempt.get("failure_reason") or "详情未取得"
            reason_label = {
                "pdf_invalid_or_corrupt": "官方 PDF 无效或损坏",
                "invalid_pdf_cooldown": "损坏 PDF 冷却期内已跳过",
                "pdf_document_unavailable": "PDF 文档暂未就绪",
                "collector_timeout": "采集器等待超时",
                "incomplete_pdf_pages": "PDF 页面不完整",
                "ocr_failure": "扫描页 OCR 未识别正文",
                "ocr_timeout": "扫描页 OCR 超时",
                "official_content_unavailable": "官方正文暂停或停止提供",
            }.get(attempt_reason, attempt_reason)
            attempt_message = last_attempt.get("message") or "本轮未取得新正文"
            _add_label_value(
                doc,
                "本轮采集：",
                f"未取得新正文（{reason_label}）：{attempt_message}",
            )
            if last_attempt.get("cooldown_until"):
                _add_label_value(
                    doc,
                    "下次自动尝试：",
                    str(last_attempt["cooldown_until"]),
                )
        _add_label_value(doc, "数据模式：", it.get("data_mode") or "live")
        _add_label_value(
            doc,
            "数据完整度：",
            (it.get("completeness") or {}).get("label") or _MISSING,
        )

        opportunity = it.get("opportunity_analysis") or {}
        _add_para(
            doc,
            "投标决策分析：" if lifecycle_stage == "机会公告" else "生命周期情报分析：",
            bold=True,
        )
        if lifecycle_stage == "机会公告":
            _add_label_value(doc, "AI/规则参与建议：", opportunity.get("decision") or "信息不足")
            _add_label_value(doc, "机会优先级：", opportunity.get("priority") or "待核验")
            _add_label_value(doc, "时间紧迫度：", opportunity.get("deadline_urgency") or "未知")
        else:
            _add_label_value(doc, "参与建议：", "不适用（该记录仅作为项目生命周期情报）")
        if opportunity.get("deadline_note"):
            _add_para(doc, f"· {opportunity['deadline_note']}", size=10)
        if opportunity.get("gaps"):
            _add_para(doc, "需补齐/核验：", bold=True, size=10)
            for gap in opportunity.get("gaps") or []:
                _add_bullet(doc, str(gap), size=10)
        if opportunity.get("recommended_actions"):
            _add_para(doc, "建议下一步：", bold=True, size=10)
            for action in opportunity.get("recommended_actions") or []:
                _add_bullet(doc, str(action), size=10)
        if opportunity.get("timeline"):
            _add_para(doc, "时间倒排：", bold=True, size=10)
            for milestone in opportunity.get("timeline") or []:
                suffix = (
                    f"（证据 {milestone.get('evidence_id')}）"
                    if milestone.get("evidence_id")
                    else ""
                )
                _add_bullet(
                    doc,
                    f"{milestone.get('milestone')}：{milestone.get('time')}{suffix}",
                    size=10,
                )
        if opportunity.get("technical_business_risks"):
            _add_para(doc, "技术/商务风险：", bold=True, size=10)
            for risk in opportunity.get("technical_business_risks") or []:
                _add_bullet(doc, str(risk), size=10)
        if opportunity.get("missing_materials"):
            _add_para(doc, "建议准备材料：", bold=True, size=10)
            for material in opportunity.get("missing_materials") or []:
                _add_bullet(doc, str(material), size=10)
        matrix = opportunity.get("qualification_matrix") or []
        if matrix:
            _add_para(doc, "资格符合性矩阵：", bold=True, size=10)
            _add_table(
                doc,
                ["条款", "原文要求", "匹配状态", "画像依据"],
                [
                    [
                        str(row.get("clause_id") or ""),
                        str(row.get("requirement") or ""),
                        str(row.get("status") or "待核验"),
                        str(row.get("profile_basis") or ""),
                    ]
                    for row in matrix
                    if isinstance(row, dict)
                ],
                col_widths_cm=[1.2, 8.0, 2.4, 4.6],
            )
        if opportunity.get("decision_evidence_ids"):
            _add_label_value(
                doc,
                "建议证据 ID：",
                "、".join(map(str, opportunity.get("decision_evidence_ids") or [])),
            )
        if opportunity.get("llm_note"):
            _add_para(doc, f"辅助研判（已校验证据）：{opportunity['llm_note']}", size=10)

        att = it.get("attachment") or {}
        _add_label_value(doc, "附件状态：", att.get("label") or "未发现公开附件")
        if att.get("note"):
            _add_para(doc, f"附件说明：{att['note']}", size=10, color=RGBColor(0x66, 0x66, 0x66))
        for a in att.get("links") or []:
            pp = doc.add_paragraph()
            r = pp.add_run("附件：")
            _set_run_font(r, size=10)
            if str(a).startswith("http"):
                _add_hyperlink(pp, str(a), "下载/查看附件")
            else:
                rr = pp.add_run(str(a))
                _set_run_font(rr, size=10)

        mb = it.get("match_basis") or {}
        _add_para(doc, "匹配依据：", bold=True)
        _add_para(doc, f"· 关键词：{mb.get('keyword', _MISSING)}", size=10)
        _add_para(doc, f"· 区域：{mb.get('region', _MISSING)}", size=10)
        _add_para(doc, f"· 时间：{mb.get('time', _MISSING)}", size=10)
        evidence = it.get("field_evidence") or {}
        evidence_labels = [
            str(value.get("source_label"))
            for value in evidence.values()
            if isinstance(value, dict) and value.get("source_label")
        ]
        if evidence_labels:
            _add_para(doc, f"字段依据：{'、'.join(dict.fromkeys(evidence_labels))}", size=9, color=RGBColor(0x66, 0x66, 0x66))
        evidence_rows = []
        for field_name, value in evidence.items():
            if not isinstance(value, dict) or not value.get("quote"):
                continue
            evidence_rows.append(
                [
                    str(value.get("evidence_id") or field_name),
                    str(value.get("source_label") or field_name),
                    str(value.get("page") or "—"),
                    str(value.get("quote") or "")[:220],
                ]
            )
        if evidence_rows:
            _add_para(doc, "关键字段原文证据：", bold=True, size=10)
            _add_table(
                doc,
                ["证据 ID", "原文标签", "PDF 页码", "原文片段"],
                evidence_rows[:18],
                col_widths_cm=[2.5, 2.4, 1.4, 9.9],
            )

        # 来源链接 — 超链接文字
        p = doc.add_paragraph()
        r = p.add_run("来源：")
        _set_run_font(r, size=11, bold=True)
        r2 = p.add_run(f"{it.get('source_display') or source_display_name(it.get('source_name'))} · ")
        _set_run_font(r2, size=11)
        url = it.get("source_url") or ""
        if str(url).startswith("http"):
            _add_hyperlink(p, str(url), "查看原始公告")
        else:
            rr = p.add_run(_MISSING)
            _set_run_font(rr, size=11)

        related = it.get("related_urls") or []
        extra_urls = [u for u in related if u and u != url]
        if extra_urls:
            _add_para(doc, "相关来源：", bold=True, size=10)
            for u in extra_urls[:5]:
                pp = doc.add_paragraph()
                if str(u).startswith("http"):
                    _add_hyperlink(pp, str(u), "查看相关公告")
                else:
                    run = pp.add_run(str(u))
                    _set_run_font(run, size=10)

        if it.get("is_update"):
            _add_para(
                doc,
                "说明：该条为已推送公告的内容更新，非全新公告。",
                bold=True,
                size=10,
            )


def _last_page(doc: Document, ctx: ReportContext, items: list[dict]) -> None:
    doc.add_page_break()
    _add_heading(doc, "九、数据处理说明与声明", 1)

    _add_heading(doc, "数据处理说明", 2)
    for line in [
        "流程：多源列表检索 → 列表条件过滤 → 详情抓取 → 详情二次过滤 → 跨源去重 → 增量判定 → 生成报告。",
        "过滤维度：关键词、区域、发布时间、招标/采购相关性。",
        "去重（入库与增量）：同批跨源合并 + 库内项目编号/内容指纹合并；主记录优先信息更完整的官方来源。",
        (
            "未去重快照：本轮每个来源采到的合格记录均写入本报告，可能重复项会保留并标注；"
            "本次不读取也不写入 DeliveryHistory。"
            if ctx.report_scope == "snapshot"
            else "增量：相对本任务历史交付记录，仅将新增或内容变化条目写入本报告。"
        ),
        "字段抽取：规则快速提取 → 关键字段缺口 AI 抽取 → 原文证据校验 → 规则结果保底；无法定位的 AI 值会被拒绝。",
    ]:
        _add_para(doc, f"· {line}")

    _add_heading(doc, "失败数据源及原因", 2)
    if ctx.sources_failed:
        for name, reason in ctx.sources_failed.items():
            _add_para(
                doc,
                f"· {source_display_name(name)}（{name}）：{reason or '未知原因'}",
            )
    else:
        _add_para(doc, "· 无失败数据源记录。")

    # 成功但零贡献
    succ = set(ctx.sources_succeeded or ctx.sources or [])
    contrib = {it.get("source_name") for it in items}
    zero = [s for s in succ if s not in contrib]
    if zero:
        _add_heading(doc, "成功但未贡献报告条目的数据源", 2)
        for s in zero:
            _add_para(
                doc,
                f"· {source_display_name(s)}：检索成功，但经过滤/去重/增量后无条目进入本报告"
                "（可能列表无命中、详情失败或均被过滤）。",
            )

    _add_label_value(
        doc,
        "抓取时间：",
        format_cn_datetime(ctx.crawl_time)
        if ctx.crawl_time
        else format_cn_datetime(ctx.generated_at),
    )
    _add_label_value(doc, "报告生成时间：", format_cn_datetime(ctx.generated_at))
    _add_label_value(doc, "执行状态：", _status_label(ctx.execution_status))

    _add_heading(doc, "数据质量与免责声明", 2)
    for line in [
        "本报告信息来源于公开网页自动抓取与规则抽取，可能存在延迟、缺漏或解析偏差。",
        "预算、截止时间、资格要求等关键商务条件请以原始公告为准。",
        "系统不绕过登录验证、不破解付费内容；登录态源不可用时自动跳过且不影响公开源。",
        "请勿将本报告作为唯一投标决策依据。",
        f"执行状态：{_status_label(ctx.execution_status)}。"
        + (
            "部分数据源未成功时，结果覆盖范围有限。"
            if ctx.execution_status == "partial"
            else ""
        ),
    ]:
        _add_para(doc, f"· {line}")


def build_word_document(ctx: ReportContext) -> Document:
    # 兼容旧调用：仅有 filtered_out_count 时拆分
    if ctx.list_filtered_out == 0 and ctx.detail_filtered_out == 0 and ctx.filtered_out_count:
        ctx.list_filtered_out = ctx.filtered_out_count
    if ctx.candidates_count == 0 and ctx.detail_success_count:
        # 粗略回退：无法精确时用 detail - detail_filtered
        ctx.candidates_count = max(
            0, ctx.detail_success_count - ctx.detail_filtered_out
        )
    if ctx.primary_count == 0:
        ctx.primary_count = ctx.final_count
    if not ctx.sources_succeeded and ctx.sources:
        ctx.sources_succeeded = list(ctx.sources)
    if not ctx.sources_requested:
        ctx.sources_requested = list(
            dict.fromkeys(
                list(ctx.sources_succeeded)
                + list((ctx.sources_failed or {}).keys())
            )
        )

    items = _prepare_items(ctx)
    doc = Document()
    # standard_business_brief tokens; Chinese font is a consistent named override.
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    pf = style.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after = Pt(6)
    pf.line_spacing = 1.10
    for name, size, color, before, after in (
        ("Heading 1", 16, "2E74B5", 16, 8),
        ("Heading 2", 13, "2E74B5", 12, 6),
        ("Heading 3", 12, "1F4D78", 8, 4),
    ):
        heading = doc.styles[name]
        heading.font.name = "Calibri"
        heading.font.size = Pt(size)
        heading.font.color.rgb = RGBColor.from_string(color)
        heading._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
        heading.paragraph_format.space_before = Pt(before)
        heading.paragraph_format.space_after = Pt(after)
        heading.paragraph_format.line_spacing = 1.10

    _add_header_footer(doc, ctx)
    _page1_cover(doc, ctx, items)
    _page2_summary(doc, ctx, items)
    _page3_overview(doc, ctx, items)
    _detail_pages(doc, ctx, items)
    _last_page(doc, ctx, items)
    return doc


def generate_report_file(
    ctx: ReportContext,
    *,
    reports_dir: Path,
) -> Path:
    """生成 Word 并写入磁盘，返回绝对路径."""
    path = build_report_filename(
        ctx.original_query,
        when=ctx.generated_at,
        reports_dir=reports_dir,
    )
    doc = build_word_document(ctx)
    path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(path))
    return path.resolve()
